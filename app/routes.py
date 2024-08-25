from flask import request, jsonify, render_template, send_file
from datetime import datetime
import logging
import colorama
from colorama import Fore, Style
from app import app
from app.init_presidio import recognize_and_anonymize_entities
from docx import Document
import openpyxl
import io

# Initialize colorama
colorama.init(autoreset=True)

# Configure logging
logging.basicConfig(filename="log.txt", level=logging.INFO)


def log_event(message, color):
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    colored_message = f"{color}[{timestamp}] [<RE-DACT>] {message}{Style.RESET_ALL}"
    logging.info(colored_message)
    print(colored_message)


def extract_text_from_word(file):
    doc = Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text), doc


def extract_text_from_excel(file):
    workbook = openpyxl.load_workbook(file)
    sheet = workbook.active
    full_text = []
    for row in sheet.iter_rows(values_only=True):
        row_text = [str(cell) if cell is not None else "" for cell in row]
        full_text.append(" ".join(row_text))
    return "\n".join(full_text), workbook


def redact_word_document(doc, redacted_text):
    paragraphs = redacted_text.split("\n")
    for i, para in enumerate(doc.paragraphs):
        if i < len(paragraphs):
            para.text = paragraphs[i]
    return doc


def redact_excel_document(workbook, redacted_text):
    sheet = workbook.active
    rows = redacted_text.split("\n")
    for i, row in enumerate(rows):
        cells = row.split(" ")
        for j, cell in enumerate(cells):
            sheet.cell(row=i + 1, column=j + 1).value = cell
    return workbook


@app.route("/")
def index():
    log_event("Web interface accessed", Fore.CYAN)
    return render_template("index.html")


@app.route("/redact", methods=["POST"])
def redact_text_route():
    if request.method == "POST":
        input_text = request.json.get("text", "")
        redaction_level = int(request.json.get("redaction_level", 100))
        log_event("Text input received", Fore.YELLOW)
        start_time = datetime.now()
        redacted_output = recognize_and_anonymize_entities(input_text, redaction_level)
        end_time = datetime.now()
        elapsed_time = (end_time - start_time).total_seconds()
        log_event(f"Redaction complete. ({elapsed_time:.2f}s)", Fore.GREEN)
        return jsonify({"redacted_text": redacted_output})


@app.route("/upload", methods=["POST"])
def upload_file():
    if "file" not in request.files:
        log_event("File upload attempted without file", Fore.RED)
        return jsonify({"error": "No file part"}), 400

    file = request.files["file"]
    if file.filename == "":
        log_event("File upload attempted with no selected file", Fore.RED)
        return jsonify({"error": "No selected file"}), 400

    file_type = None
    content = None

    if file and file.filename.endswith(".txt"):
        content = file.read().decode("utf-8")
        file_type = "text"
    elif file and file.filename.endswith(".docx"):
        content, doc = extract_text_from_word(file)
        file_type = "word"
    elif file and file.filename.endswith(".xlsx"):
        content, workbook = extract_text_from_excel(file)
        file_type = "excel"
    else:
        log_event("Invalid file format attempted", Fore.RED)
        return (
            jsonify(
                {"error": "Invalid file format. Supported formats: .txt, .docx, .xlsx"}
            ),
            400,
        )

    log_event(f"File '{file.filename}' uploaded", Fore.BLUE)
    original_text = content

    # Store the original file for later redaction
    if file_type == "word":
        app.config["original_file"] = doc
    elif file_type == "excel":
        app.config["original_file"] = workbook
    else:
        app.config["original_file"] = None

    app.config["file_type"] = file_type
    app.config["original_text"] = original_text

    return jsonify({"original_text": original_text})


@app.route("/download", methods=["POST"])
def download_file():
    redaction_level = int(request.json.get("redaction_level", 100))
    original_text = app.config.get("original_text", "")
    file_type = app.config.get("file_type", "text")
    start_time = datetime.now()

    # Perform redaction
    redacted_text = recognize_and_anonymize_entities(original_text, redaction_level)
    end_time = datetime.now()
    elapsed_time = (end_time - start_time).total_seconds()
    log_event(f"Redaction complete. ({elapsed_time:.2f}s)", Fore.GREEN)

    # Handle different file types
    if file_type == "word":
        doc = app.config["original_file"]
        redacted_doc = redact_word_document(doc, redacted_text)
        buffer = io.BytesIO()
        redacted_doc.save(buffer)
        buffer.seek(0)
        return send_file(
            buffer,
            as_attachment=True,
            download_name="redacted_document.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    elif file_type == "excel":
        workbook = app.config["original_file"]
        redacted_workbook = redact_excel_document(workbook, redacted_text)
        buffer = io.BytesIO()
        redacted_workbook.save(buffer)
        buffer.seek(0)
        return send_file(
            buffer,
            as_attachment=True,
            download_name="redacted_document.xlsx",
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

    else:  # Default to text
        buffer = io.StringIO(redacted_text)
        buffer.seek(0)
        return send_file(
            io.BytesIO(buffer.getvalue().encode()),
            as_attachment=True,
            download_name="redacted_text.txt",
            mimetype="text/plain",
        )
